import os
import re
import json
import hashlib
import time
import random
import gc
from datetime import datetime
from flask import Flask, render_template, request, jsonify
from dotenv import load_dotenv
from google import genai
from google.genai import types
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np

load_dotenv()

app = Flask(__name__)
app.secret_key = os.urandom(24)

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# ── Gemini client ──────────────────────────────────────────────────────────────
gemini_client = genai.Client(api_key=os.environ.get("GEMINI_API_KEY"))
GEMINI_MODEL = "gemini-2.5-flash"

# ── Halacha system prompt (kept small — document chunks go in user message) ──
SYSTEM_PROMPT = """אתה פוסק הלכה מומחה המתמחה אך ורק בספרי "ילקוט יוסף" של הרב עובדיה יוסף זצ"ל והרב יצחק יוסף שליט"א.

כללים מחייבים:
1. עליך לענות על שאלות אך ורק על סמך קטעי המסמך המצורפים. אין להשתמש בידע חיצוני, אין להוסיף לוגיקה משלך, ואין לשנות את משמעות הטקסט. "רק מה שכתוב".
2. תמיד ציין את המקור המדויק (פרק, סימן, סעיף, הלכה, עמוד וכו') מתוך הקטעים לצד תשובתך.
3. אם התשובה לא נמצאת בקטעים המצורפים, השב בדיוק: "התשובה לא נמצאת בקובץ המקור".
4. ענה בעברית בלבד.
5. השתמש בשפה הלכתית ברורה ומדויקת.
6. אם יש מחלוקת פוסקים המוזכרת בקטעים, ציין אותה כפי שמופיעה בטקסט.
7. אל תמציא מקורות או מראי מקומות שלא מופיעים בקטעים.
8. כאשר אתה מצטט, ציטוט מדויק ככל האפשר מהטקסט המקורי.

כללי עיצוב התשובה:
- השתמש ב-**כוכביות כפולות** להדגשת מונחים הלכתיים חשובים ומילות מפתח.
- חלק את התשובה לפסקאות ברורות ומסודרות.
- כאשר יש מספר דינים או פרטים, השתמש ברשימה ממוספרת (1. 2. 3.).
- בסוף התשובה, תמיד הוסף שורה ריקה ואז שורה שמתחילה בדיוק ב-"מקור:" ואחריה הציון המדויק של המקור מתוך הקטעים (לדוגמה: "מקור: סימן שב, סעיף א, עמוד 45")."""


# ══════════════════════════════════════════════════════════════════════════════
#  Feature #16: Hebrew stop-word removal
# ══════════════════════════════════════════════════════════════════════════════
HEBREW_STOP_WORDS = {
    "של", "את", "הוא", "היא", "אני", "על", "עם", "כי", "לא", "אם",
    "גם", "או", "כל", "אבל", "הם", "הן", "אנחנו", "יש", "אין",
    "רק", "עד", "מן", "בין", "כן", "אלא", "כמו", "לפי", "אותו",
    "אותה", "שלא", "אשר", "כאשר", "היה", "היו", "היתה", "יהיה",
    "להיות", "זה", "זו", "זאת", "אלה", "אלו", "כך", "כבר", "עוד",
    "מאד", "ביותר", "אחר", "אחרי", "לפני", "אצל", "בלי", "למה",
    "מה", "איך", "מי", "כמה", "מתי", "היכן", "איפה", "שם", "פה",
    "הנה", "בו", "בה", "בהם", "שהוא", "שהיא", "שהם", "ואם", "וכן",
}


def remove_stop_words(text):
    """Remove Hebrew stop words from text for better search quality."""
    words = re.findall(r'[\u0590-\u05FF\w]+', text)
    return ' '.join(w for w in words if w not in HEBREW_STOP_WORDS)


# ══════════════════════════════════════════════════════════════════════════════
#  Feature #18: Auto-detect Halacha topic
# ══════════════════════════════════════════════════════════════════════════════
HALACHA_TOPICS = {
    "שבת": {"keywords": ["שבת", "מוצאי שבת", "ערב שבת", "שביתה", "מלאכה", "מלאכות", "מוקצה", "הבדלה", "קידוש"], "icon": "🕯️"},
    "ברכות": {"keywords": ["ברכה", "ברכות", "ברכת", "שהכל", "העץ", "האדמה", "מזונות", "המוציא", "ברכת המזון"], "icon": "🙏"},
    "תפילה": {"keywords": ["תפילה", "תפילת", "שמונה עשרה", "עמידה", "שחרית", "מנחה", "ערבית", "מעריב", "קריאת שמע"], "icon": "📖"},
    "כשרות": {"keywords": ["כשר", "כשרות", "טרף", "בשר", "חלב", "בשרי", "חלבי", "הכשר", "שחיטה", "תערובות"], "icon": "🍽️"},
    "חגים ומועדים": {"keywords": ["פסח", "סוכות", "שבועות", "ראש השנה", "יום כיפור", "חנוכה", "פורים", "סוכה", "מצה", "שופר", "חג"], "icon": "📅"},
    "נידה וטהרה": {"keywords": ["נידה", "טהרה", "מקוה", "טבילה", "טהרת המשפחה"], "icon": "💧"},
    "שמיטה": {"keywords": ["שמיטה", "שביעית", "היתר מכירה", "אוצר בית דין"], "icon": "🌱"},
    "אבלות": {"keywords": ["אבלות", "אבל", "שבעה", "שלושים", "הספד", "קריעה", "ניחום אבלים"], "icon": "🕊️"},
    "נישואין": {"keywords": ["נישואין", "חתונה", "קידושין", "כתובה", "חופה"], "icon": "💍"},
    "מזוזה ותפילין": {"keywords": ["מזוזה", "תפילין", "ציצית", "מזוזות"], "icon": "📜"},
    "צדקה וחסד": {"keywords": ["צדקה", "מעשר", "חסד", "גמילות חסדים", "נתינה"], "icon": "🤲"},
    "בין אדם לחברו": {"keywords": ["לשון הרע", "רכילות", "הלבנת פנים", "גניבת דעת", "אונאה"], "icon": "🤝"},
}


def detect_topic(question):
    """Detect the halacha topic from a question. Returns (topic_name, icon) or (None, None)."""
    question_lower = question.strip()
    best_topic = None
    best_count = 0

    for topic, info in HALACHA_TOPICS.items():
        count = sum(1 for kw in info["keywords"] if kw in question_lower)
        if count > best_count:
            best_count = count
            best_topic = topic

    if best_topic and best_count > 0:
        return best_topic, HALACHA_TOPICS[best_topic]["icon"]
    return None, None


# ══════════════════════════════════════════════════════════════════════════════
#  Feature #14: Query Cache (JSON file)
# ══════════════════════════════════════════════════════════════════════════════
CACHE_FILE = os.path.join(BASE_DIR, "query_cache.json")
query_cache = {}


def load_cache():
    """Load query cache from disk."""
    global query_cache
    if os.path.exists(CACHE_FILE):
        try:
            with open(CACHE_FILE, 'r', encoding='utf-8') as f:
                query_cache = json.load(f)
            print(f"[✓] טעינת מטמון: {len(query_cache)} שאילתות")
        except Exception:
            query_cache = {}


def save_cache():
    """Persist query cache to disk."""
    try:
        with open(CACHE_FILE, 'w', encoding='utf-8') as f:
            json.dump(query_cache, f, ensure_ascii=False, indent=2)
    except Exception:
        pass


def get_cache_key(question):
    """Generate a deterministic cache key for a question."""
    normalized = re.sub(r'\s+', ' ', question.strip())
    return hashlib.md5(normalized.encode('utf-8')).hexdigest()


# ══════════════════════════════════════════════════════════════════════════════
#  Feature #19: Search Logging
# ══════════════════════════════════════════════════════════════════════════════
LOG_FILE = os.path.join(BASE_DIR, "search_logs.json")


def log_search(question, topic, found, chunks_used, cached):
    """Append a search event to search_logs.json."""
    entry = {
        "timestamp": datetime.now().isoformat(),
        "question": question,
        "topic": topic,
        "answer_found": found,
        "chunks_used": chunks_used,
        "from_cache": cached,
    }
    try:
        logs = []
        if os.path.exists(LOG_FILE):
            with open(LOG_FILE, 'r', encoding='utf-8') as f:
                logs = json.load(f)
        logs.append(entry)
        with open(LOG_FILE, 'w', encoding='utf-8') as f:
            json.dump(logs, f, ensure_ascii=False, indent=2)
    except Exception:
        pass


# ── Text extraction helpers ───────────────────────────────────────────────────
def extract_text_from_pdf(filepath):
    """Extract text from a PDF file, preserving page numbers."""
    reader = PdfReader(filepath)
    text_parts = []
    total_pages = len(reader.pages)
    for i, page in enumerate(reader.pages, 1):
        # Feature #9: Terminal progress
        if i % 50 == 0 or i == total_pages:
            pct = int(i / total_pages * 100)
            bar = "█" * (pct // 5) + "░" * (20 - pct // 5)
            print(f"\r    [{bar}] {pct}% ({i}/{total_pages} עמודים)", end="", flush=True)
        page_text = page.extract_text()
        if page_text:
            text_parts.append(f"--- עמוד {i} ---\n{page_text}")
    print()  # newline after progress bar
    return "\n\n".join(text_parts)


def extract_text_from_docx(filepath):
    """Extract text from a DOCX file."""
    doc = DocxDocument(filepath)
    text_parts = []
    total_paras = len(doc.paragraphs)
    for i, para in enumerate(doc.paragraphs, 1):
        # Feature #9: Terminal progress
        if i % 500 == 0 or i == total_paras:
            pct = int(i / total_paras * 100)
            bar = "█" * (pct // 5) + "░" * (20 - pct // 5)
            print(f"\r    [{bar}] {pct}% ({i}/{total_paras} פסקאות)", end="", flush=True)
        if para.text.strip():
            text_parts.append(para.text)
    print()  # newline after progress bar
    return "\n".join(text_parts)


# ══════════════════════════════════════════════════════════════════════════════
#  RAG Index — local TF-IDF retrieval (zero API calls)
# ══════════════════════════════════════════════════════════════════════════════

class RAGIndex:
    """Lightweight local search index using TF-IDF + cosine similarity.

    - Chunks the document into meaningful sections (~500-2000 chars)
    - Feature #15: Chunk overlap for better context continuity
    - Feature #17: Character n-gram analyzer for fuzzy/typo tolerance
    - Builds a TF-IDF matrix entirely offline (no API calls)
    - At query time, vectorises the question and returns the top-K chunks
    """

    # Regex patterns that mark section boundaries in Yalkut Yosef text
    SECTION_PATTERNS = re.compile(
        r"(?:^|\n)"                       # start of line
        r"(?:"
        r"סימן\s+[א-ת\"\']+|"            # סימן א', סימן ב', ...
        r"פרק\s+[א-ת\"\'0-9]+|"          # פרק א', ...
        r"הלכות\s+\S+|"                   # הלכות שבת, ...
        r"סעיף\s+[א-ת\"\'0-9]+|"         # סעיף א', ...
        r"---\s*עמוד\s+\d+\s*---"         # --- עמוד 5 --- (PDF page markers)
        r")",
        re.MULTILINE,
    )

    def __init__(self, min_chunk_chars=200, max_chunk_chars=2000, overlap_chars=150):
        self.min_chunk = min_chunk_chars
        self.max_chunk = max_chunk_chars
        self.overlap_chars = overlap_chars  # Feature #15
        self.chunks: list[str] = []
        self.word_vectorizer: TfidfVectorizer | None = None
        self.char_vectorizer: TfidfVectorizer | None = None  # Feature #17
        self.word_matrix = None
        self.char_matrix = None

    # ── 1. Chunking (with overlap) ────────────────────────────────────────
    def _split_into_chunks(self, text: str) -> list[str]:
        """Split text into chunks, trying to respect section boundaries.
        Feature #15: Adds overlap between chunks for better context continuity."""
        # First pass: split on blank-line boundaries
        raw_paragraphs = re.split(r"\n{2,}", text)

        chunks: list[str] = []
        current = ""

        for para in raw_paragraphs:
            para = para.strip()
            if not para:
                continue

            # Check if this paragraph starts a new section
            is_section_start = bool(self.SECTION_PATTERNS.match(para))

            # If adding this paragraph would exceed max, or it's a new section
            # and current is already substantial, flush current chunk
            if current and (
                len(current) + len(para) + 1 > self.max_chunk
                or (is_section_start and len(current) >= self.min_chunk)
            ):
                chunks.append(current)
                current = para
            else:
                current = current + "\n" + para if current else para

        if current and len(current.strip()) > 30:
            chunks.append(current)

        # Second pass: merge tiny chunks with their neighbour
        merged: list[str] = []
        for chunk in chunks:
            if merged and len(merged[-1]) < self.min_chunk:
                merged[-1] = merged[-1] + "\n" + chunk
            else:
                merged.append(chunk)

        # Feature #15: Add overlap — carry tail of previous chunk into next
        if self.overlap_chars > 0 and len(merged) > 1:
            overlapped = [merged[0]]
            for i in range(1, len(merged)):
                prev_tail = merged[i - 1][-self.overlap_chars:]
                overlapped.append(prev_tail + "\n" + merged[i])
            return overlapped

        return merged

    # ── 2. Build index ───────────────────────────────────────────────────
    def build(self, text: str):
        """Chunk the text and build the TF-IDF index."""
        t0 = time.time()

        # Feature #9: Terminal progress
        print("    [שלב 1/3] פיצול לקטעים...")
        self.chunks = self._split_into_chunks(text)
        print(f"    קטעים שנוצרו: {len(self.chunks):,}")

        if not self.chunks:
            print("[!] אין קטעים – אינדקס לא נוצר")
            return

        # Feature #9: Terminal progress
        print("    [שלב 2/3] בניית אינדקס מילים (TF-IDF)...")

        # Word-level TF-IDF (memory-optimized: max_features + float32)
        self.word_vectorizer = TfidfVectorizer(
            analyzer="word",
            token_pattern=r"[\u0590-\u05FF\w]{2,}",  # Hebrew + ASCII words, 2+ chars
            sublinear_tf=True,       # log(1+tf) — BM25-like dampening
            min_df=2,                # ignore ultra-rare terms
            max_df=0.85,             # ignore terms in >85% of chunks
            ngram_range=(1, 2),      # unigrams + bigrams
            max_features=5000,       # limit vocabulary to save RAM
            dtype=np.float32,        # half the memory vs float64
        )
        self.word_matrix = self.word_vectorizer.fit_transform(self.chunks)

        vocab_size = len(self.word_vectorizer.vocabulary_)
        print(f"    גודל מילון מילים: {vocab_size:,} מונחים")

        # Feature #17: Character n-gram TF-IDF for fuzzy matching / typo tolerance
        print("    [שלב 3/3] בניית אינדקס תווים (fuzzy matching)...")
        self.char_vectorizer = TfidfVectorizer(
            analyzer="char_wb",
            ngram_range=(3, 5),      # character 3-grams to 5-grams
            min_df=3,
            max_df=0.9,
            sublinear_tf=True,
            max_features=8000,       # limit char n-gram vocabulary to save RAM
            dtype=np.float32,        # half the memory vs float64
        )
        self.char_matrix = self.char_vectorizer.fit_transform(self.chunks)

        char_vocab_size = len(self.char_vectorizer.vocabulary_)
        print(f"    גודל מילון תווים: {char_vocab_size:,} n-grams")

        elapsed = time.time() - t0
        print(f"    זמן בניית אינדקס: {elapsed:.1f} שניות")

    # ── 3. Search (hybrid word + char n-gram) ────────────────────────────
    def search(self, query: str, top_k: int = 10) -> list[dict]:
        """Return the top-K most relevant chunks for the query.
        Feature #16: Applies Hebrew stop-word removal.
        Feature #17: Combines word and char n-gram scores for fuzzy matching."""
        if self.word_vectorizer is None or self.word_matrix is None:
            return []

        # Feature #16: Remove stop words from query
        clean_query = remove_stop_words(query)

        # Word-level scores
        word_vec = self.word_vectorizer.transform([clean_query])
        word_scores = cosine_similarity(word_vec, self.word_matrix).flatten()

        # Feature #17: Character n-gram scores for fuzzy matching
        char_scores = np.zeros_like(word_scores)
        if self.char_vectorizer is not None and self.char_matrix is not None:
            char_vec = self.char_vectorizer.transform([query])  # use original query for char n-grams
            char_scores = cosine_similarity(char_vec, self.char_matrix).flatten()

        # Combine: 70% word + 30% char n-gram
        combined_scores = 0.7 * word_scores + 0.3 * char_scores

        # Get top-K indices, sorted by score descending
        top_indices = np.argsort(combined_scores)[::-1][:top_k]

        results = []
        for idx in top_indices:
            score = float(combined_scores[idx])
            if score < 0.01:
                break  # below noise floor — skip
            results.append({
                "chunk_id": int(idx),
                "score": round(score, 4),
                "text": self.chunks[idx],
            })

        return results


# ══════════════════════════════════════════════════════════════════════════════
#  Feature #20: Enhanced source parser
# ══════════════════════════════════════════════════════════════════════════════
SOURCE_PATTERN = re.compile(
    r'(?:כרך|חלק)\s+[א-ת\d"\']+|'
    r'סימן\s+[א-ת\d"\']+|'
    r'סעיף\s+[א-ת\d"\']+|'
    r'הלכה\s+[א-ת\d"\']+|'
    r'עמוד\s+\d+|'
    r'עמ\'\s*\d+|'
    r'דף\s+[א-ת\d"\']+',
    re.MULTILINE,
)


def parse_source_structure(source_text):
    """Parse a source citation into structured parts (כרך, סימן, סעיף, etc.)."""
    parts = SOURCE_PATTERN.findall(source_text)
    if parts:
        return [p.strip() for p in parts]
    return [source_text.strip()] if source_text.strip() else []


# ══════════════════════════════════════════════════════════════════════════════
#  Pre-load document + build RAG index at startup
# ══════════════════════════════════════════════════════════════════════════════

DOCUMENT_TEXT = ""
DOCUMENT_FILENAME = ""
DOCUMENT_CHAR_COUNT = 0
DOCUMENT_PAGE_COUNT = None
DOCUMENT_CHUNK_COUNT = 0

rag_index = RAGIndex(min_chunk_chars=200, max_chunk_chars=2000, overlap_chars=150)


def load_yalkut_yosef():
    """Auto-load the Yalkut Yosef document and build the RAG index."""
    global DOCUMENT_TEXT, DOCUMENT_FILENAME, DOCUMENT_CHAR_COUNT
    global DOCUMENT_PAGE_COUNT, DOCUMENT_CHUNK_COUNT

    docx_path = os.path.join(BASE_DIR, "ילקוט-יוסף-קיצור-שולחן-ערוך.docx")
    pdf_path = os.path.join(BASE_DIR, "ילקוט-יוסף-קיצור-שולחן-ערוך.pdf")

    if os.path.exists(docx_path):
        print(f"\n{'='*60}")
        print(f"[*] טוען מסמך DOCX: {os.path.basename(docx_path)}")
        print(f"{'='*60}")
        DOCUMENT_TEXT = extract_text_from_docx(docx_path)
        DOCUMENT_FILENAME = "ילקוט-יוסף-קיצור-שולחן-ערוך.docx"
    elif os.path.exists(pdf_path):
        print(f"\n{'='*60}")
        print(f"[*] טוען מסמך PDF: {os.path.basename(pdf_path)}")
        print(f"{'='*60}")
        DOCUMENT_TEXT = extract_text_from_pdf(pdf_path)
        DOCUMENT_FILENAME = "ילקוט-יוסף-קיצור-שולחן-ערוך.pdf"
    else:
        print("[!] שגיאה: לא נמצא קובץ ילקוט יוסף בתיקיית הפרויקט!")
        return

    DOCUMENT_CHAR_COUNT = len(DOCUMENT_TEXT)
    DOCUMENT_PAGE_COUNT = (
        DOCUMENT_TEXT.count("--- עמוד") if "--- עמוד" in DOCUMENT_TEXT else None
    )

    print(f"[✓] מסמך נטען בהצלחה: {DOCUMENT_FILENAME}")
    print(f"    תווים: {DOCUMENT_CHAR_COUNT:,}")
    if DOCUMENT_PAGE_COUNT:
        print(f"    עמודים: {DOCUMENT_PAGE_COUNT}")

    # Build the RAG search index
    print("[*] בונה אינדקס חיפוש (TF-IDF + char n-grams)...")
    rag_index.build(DOCUMENT_TEXT)
    DOCUMENT_CHUNK_COUNT = len(rag_index.chunks)
    print(f"[✓] אינדקס מוכן — {DOCUMENT_CHUNK_COUNT:,} קטעים")
    print(f"{'='*60}\n")


# ══════════════════════════════════════════════════════════════════════════════
#  Feature #11: Daily Halacha — smart snippet extraction
# ══════════════════════════════════════════════════════════════════════════════

# Regex for lines that open a proper halachic section/paragraph
_CLEAN_START_RE = re.compile(
    r'^(?:'
    r'סעיף\s+[א-ת\d"\']+|'        # סעיף א'
    r'סימן\s+[א-ת\d"\']+|'        # סימן שב
    r'הלכה\s+[א-ת\d"\']+|'        # הלכה ג'
    r'פרק\s+[א-ת\d"\']+|'         # פרק א'
    r'הלכות\s+\S+|'               # הלכות שבת
    r'[א-ת]'                       # any line starting with a Hebrew letter
    r')'
)

# Sentence-ending punctuation used in halachic texts
_SENTENCE_END_RE = re.compile(r'[.׃:]\s*$')

# Regex for extracting topic context (הלכות שבת, הלכות ברכות, etc.)
_TOPIC_RE = re.compile(r'^הלכות\s+(\S+(?:\s+\S+)?)')
# Regex for extracting siman context (סימן שב, סימן א', etc.)
_SIMAN_RE = re.compile(r'^סימן\s+([א-ת\d"\']+)')

daily_snippets: list[dict] = []


def build_daily_snippets():
    """Build a curated list of complete, self-contained halachic paragraphs
    suitable for the daily-halacha feature.  Each snippet dict has:
      - text: the snippet content (100-600 chars, starts/ends cleanly)
      - topic: the broader halachic topic (e.g. 'הלכות שבת')
      - siman: the exact siman (e.g. 'סימן שב')
    """
    global daily_snippets
    if not DOCUMENT_TEXT:
        return

    # Split the full document on blank lines → natural paragraphs
    raw_paragraphs = re.split(r'\n{2,}', DOCUMENT_TEXT)

    candidates: list[dict] = []
    last_topic = ""
    last_siman = ""

    for para in raw_paragraphs:
        para = para.strip()
        if not para:
            continue

        # Track topic and siman context as we scan through the document
        topic_match = _TOPIC_RE.match(para)
        if topic_match:
            last_topic = "הלכות " + topic_match.group(1)

        siman_match = _SIMAN_RE.match(para)
        if siman_match:
            last_siman = "סימן " + siman_match.group(1)

        if len(para) < 80:
            continue

        # Must start with a recognisable Hebrew opening
        if not _CLEAN_START_RE.match(para):
            continue

        # Strip page markers that bleed in from PDF extraction
        cleaned = re.sub(r'---\s*עמוד\s+\d+\s*---', '', para).strip()
        if len(cleaned) < 80:
            continue

        snippet_text = None

        # If the paragraph is short enough, take it whole
        if len(cleaned) <= 600:
            snippet_text = cleaned
        else:
            # Otherwise, trim to a sentence boundary within 600 chars
            window = cleaned[:600]
            # Search backwards for sentence-ending punctuation
            best_cut = -1
            for m in _SENTENCE_END_RE.finditer(window):
                best_cut = m.end()
            if best_cut >= 100:
                snippet_text = window[:best_cut].strip()
            else:
                # Fallback: cut at last full stop / colon
                for ch in ('.', ':', '׃'):
                    idx = window.rfind(ch)
                    if idx >= 100:
                        snippet_text = window[:idx + 1].strip()
                        break

        if snippet_text:
            candidates.append({
                "text": snippet_text,
                "topic": last_topic,
                "siman": last_siman,
            })

    daily_snippets = candidates
    print(f"[✓] הלכה יומית: {len(daily_snippets):,} קטעים ראויים")


# Load once at startup
load_yalkut_yosef()
build_daily_snippets()
load_cache()

# ── Free RAM: delete the massive raw text now that index + snippets are built ──
DOCUMENT_TEXT = ""
gc.collect()
print("[✓] זיכרון שוחרר: מחרוזת המסמך הגולמית נמחקה")


# ── Routes ────────────────────────────────────────────────────────────────────
@app.route("/")
def index():
    return render_template("index.html")


@app.route("/status", methods=["GET"])
def status():
    """Return info about the pre-loaded document so the frontend knows it is ready."""
    return jsonify({
        "ready": DOCUMENT_CHUNK_COUNT > 0,
        "filename": DOCUMENT_FILENAME,
        "char_count": DOCUMENT_CHAR_COUNT,
        "page_count": DOCUMENT_PAGE_COUNT,
        "chunk_count": DOCUMENT_CHUNK_COUNT,
    })


@app.route("/ask", methods=["POST"])
def ask_question():
    data = request.get_json()
    if not data:
        return jsonify({"error": "לא התקבלו נתונים", "error_type": "input"}), 400

    question = data.get("question", "").strip()

    if not question:
        return jsonify({"error": "נא להזין שאלה", "error_type": "input"}), 400

    if DOCUMENT_CHUNK_COUNT == 0:
        return jsonify({"error": "לא נטען מסמך מקור. ודא שקובץ ילקוט יוסף נמצא בתיקיית הפרויקט.", "error_type": "system"}), 500

    # Feature #18: Detect topic
    topic_name, topic_icon = detect_topic(question)

    # Feature #14: Check cache first
    cache_key = get_cache_key(question)
    if cache_key in query_cache:
        cached = query_cache[cache_key]
        log_search(question, topic_name, cached.get("answer_found", True), cached.get("chunks_used", 0), True)
        return jsonify({
            "success": True,
            "answer": cached["answer"],
            "source_file": cached.get("source_file", DOCUMENT_FILENAME),
            "tokens_used": cached.get("tokens_used", {}),
            "chunks_used": cached.get("chunks_used", 0),
            "from_cache": True,
            "topic": topic_name,
            "topic_icon": topic_icon,
            "source_parts": cached.get("source_parts", []),
        })

    # ── Step 1: Retrieve relevant chunks locally (no API call) ────────────
    retrieved = rag_index.search(question, top_k=10)

    if not retrieved:
        log_search(question, topic_name, False, 0, False)
        return jsonify({
            "success": True,
            "answer": "התשובה לא נמצאת בקובץ המקור",
            "source_file": DOCUMENT_FILENAME,
            "tokens_used": {},
            "topic": topic_name,
            "topic_icon": topic_icon,
            "source_parts": [],
        })

    # ── Step 2: Build a compact context from retrieved chunks ─────────────
    context_parts = []
    for i, r in enumerate(retrieved, 1):
        context_parts.append(f"[קטע {i} | רלוונטיות: {r['score']}]\n{r['text']}")

    context_block = "\n\n---\n\n".join(context_parts)

    user_message = (
        "להלן קטעים רלוונטיים מתוך ספרי ילקוט יוסף שנמצאו עבור השאלה:\n\n"
        "--- תחילת קטעים ---\n"
        f"{context_block}\n"
        "--- סוף קטעים ---\n\n"
        f"שאלה: {question}\n\n"
        "ענה אך ורק על סמך הקטעים שלמעלה. "
        "ציין מקור מדויק (סימן, סעיף, הלכה, עמוד). "
        'אם התשובה לא נמצאת בקטעים, השב: "התשובה לא נמצאת בקובץ המקור".'
    )

    # ── Step 3: Send only the small context to Gemini ─────────────────────
    try:
        response = gemini_client.models.generate_content(
            model=GEMINI_MODEL,
            contents=user_message,
            config=types.GenerateContentConfig(
                system_instruction=SYSTEM_PROMPT,
                max_output_tokens=4096,
            ),
        )

        answer = response.text

        # Extract token usage
        tokens_used = {}
        usage = getattr(response, "usage_metadata", None)
        if usage:
            tokens_used = {
                "input": getattr(usage, "prompt_token_count", 0),
                "output": getattr(usage, "candidates_token_count", 0),
            }

        # Feature #20: Parse source structure from answer
        source_parts = parse_source_structure(answer)

        answer_found = "התשובה לא נמצאת בקובץ המקור" not in answer

        # Feature #14: Cache the result
        query_cache[cache_key] = {
            "answer": answer,
            "source_file": DOCUMENT_FILENAME,
            "tokens_used": tokens_used,
            "chunks_used": len(retrieved),
            "source_parts": source_parts,
            "answer_found": answer_found,
        }
        save_cache()

        # Feature #19: Log search
        log_search(question, topic_name, answer_found, len(retrieved), False)

        return jsonify({
            "success": True,
            "answer": answer,
            "source_file": DOCUMENT_FILENAME,
            "tokens_used": tokens_used,
            "chunks_used": len(retrieved),
            "from_cache": False,
            "topic": topic_name,
            "topic_icon": topic_icon,
            "source_parts": source_parts,
        })

    except Exception as e:
        error_msg = str(e)
        # Feature #8: Enhanced error classification
        if "API_KEY" in error_msg or "api_key" in error_msg or "401" in error_msg:
            return jsonify({"error": "מפתח API לא תקין. בדוק את הגדרות GEMINI_API_KEY", "error_type": "auth"}), 401
        elif "429" in error_msg or "quota" in error_msg.lower() or "rate" in error_msg.lower():
            return jsonify({"error": "חריגה ממגבלת בקשות. נסה שוב בעוד דקה", "error_type": "rate_limit"}), 429
        elif "block" in error_msg.lower() or "safety" in error_msg.lower():
            return jsonify({"error": "התוכן נחסם על ידי מסנני הבטיחות. נסה לנסח מחדש את השאלה.", "error_type": "safety"}), 400
        elif "timeout" in error_msg.lower() or "deadline" in error_msg.lower():
            return jsonify({"error": "הבקשה ארכה יותר מדי זמן. נסה שוב.", "error_type": "timeout"}), 504
        elif "connect" in error_msg.lower() or "network" in error_msg.lower():
            return jsonify({"error": "שגיאת חיבור לשרת. בדוק את חיבור האינטרנט.", "error_type": "network"}), 503
        else:
            return jsonify({"error": f"שגיאה: {error_msg}", "error_type": "unknown"}), 500


@app.route("/daily", methods=["GET"])
def daily_halacha():
    """Return a complete, clean halachic paragraph with topic/siman context."""
    if not daily_snippets:
        return jsonify({"snippet": None, "topic": "", "siman": ""})

    # Deterministic seed from today's date → same snippet all day
    today_seed = int(datetime.now().strftime("%Y%m%d"))
    rng = random.Random(today_seed)
    entry = rng.choice(daily_snippets)

    return jsonify({
        "snippet": entry["text"],
        "topic": entry.get("topic", ""),
        "siman": entry.get("siman", ""),
    })


# ══════════════════════════════════════════════════════════════════════════════
#  Feature #10: Feedback endpoint
# ══════════════════════════════════════════════════════════════════════════════
FEEDBACK_FILE = os.path.join(BASE_DIR, "feedback.json")


@app.route("/feedback", methods=["POST"])
def submit_feedback():
    """Record thumbs up/down feedback for a Q&A pair."""
    data = request.get_json()
    if not data:
        return jsonify({"error": "לא התקבלו נתונים"}), 400

    entry = {
        "timestamp": datetime.now().isoformat(),
        "question": data.get("question", ""),
        "rating": data.get("rating", ""),  # "up" or "down"
    }

    try:
        feedback = []
        if os.path.exists(FEEDBACK_FILE):
            with open(FEEDBACK_FILE, 'r', encoding='utf-8') as f:
                feedback = json.load(f)
        feedback.append(entry)
        with open(FEEDBACK_FILE, 'w', encoding='utf-8') as f:
            json.dump(feedback, f, ensure_ascii=False, indent=2)
        return jsonify({"success": True})
    except Exception:
        return jsonify({"error": "שגיאה בשמירת המשוב"}), 500


if __name__ == "__main__":
    app.run(debug=True, host="0.0.0.0", port=5000)
