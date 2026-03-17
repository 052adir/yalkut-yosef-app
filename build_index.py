"""
build_index.py — Pre-compute the TF-IDF index and chunks offline.

Run this script locally BEFORE deploying to Render:
    python build_index.py

It reads the Yalkut Yosef .docx, splits it into chunks, trains the
TF-IDF vectorizers, builds the daily-halacha snippets, and saves
everything to the data/ folder as lightweight binary files.
"""

import os
import re
import time
import json
import random
import joblib
import numpy as np
from docx import Document as DocxDocument
from PyPDF2 import PdfReader
from sklearn.feature_extraction.text import TfidfVectorizer

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DATA_DIR = os.path.join(BASE_DIR, "data")

# ── Hebrew stop words (same as app.py) ────────────────────────────────────────
HEBREW_STOP_WORDS = {
    "של", "את", "הוא", "היא", "אני", "על", "עם", "כי", "לא", "אם",
    "גם", "או", "כל", "אבל", "הם", "הן", "אנחנו", "יש", "אין",
    "רק", "עד", "מן", "בין", "כן", "אלא", "כמו", "לפי", "אותו",
    "אותה", "שלא", "אשר", "כאשר", "היה", "היו", "היתה", "יהיה",
    "להיות", "זה", "זו", "זאת", "אלה", "אלו", "כך", "כבר", "עוד",
    "מאד", "ביותר", "אחר", "אחרי", "לפני", "אצל", "בלי", "למה",
    "מה", "איך", "מי", "כמה", "מתי", "היכן", "איפה", "שם", "פה",
    "הנה", "בו", "בה", "בהם", "שהוא", "שהיא", "שהם", "ואם", "וכן",
}

# ── Section / chunking patterns ───────────────────────────────────────────────
SECTION_PATTERNS = re.compile(
    r"(?:^|\n)"
    r"(?:"
    r"סימן\s+[א-ת\"\']+|"
    r"פרק\s+[א-ת\"\'0-9]+|"
    r"הלכות\s+\S+|"
    r"סעיף\s+[א-ת\"\'0-9]+|"
    r"---\s*עמוד\s+\d+\s*---"
    r")",
    re.MULTILINE,
)

# ── Daily-snippet patterns ────────────────────────────────────────────────────
_CLEAN_START_RE = re.compile(
    r'^(?:'
    r'סעיף\s+[א-ת\d"\']+|'
    r'סימן\s+[א-ת\d"\']+|'
    r'הלכה\s+[א-ת\d"\']+|'
    r'פרק\s+[א-ת\d"\']+|'
    r'הלכות\s+\S+|'
    r'[א-ת]'
    r')'
)
_SENTENCE_END_RE = re.compile(r'[.׃:]\s*$')
_TOPIC_RE = re.compile(r'^הלכות\s+(\S+(?:\s+\S+)?)')
_SIMAN_RE = re.compile(r'^סימן\s+([א-ת\d"\']+)')


# ══════════════════════════════════════════════════════════════════════════════
#  Text extraction
# ══════════════════════════════════════════════════════════════════════════════

def extract_text_from_docx(filepath):
    doc = DocxDocument(filepath)
    text_parts = []
    total_paras = len(doc.paragraphs)
    for i, para in enumerate(doc.paragraphs, 1):
        if i % 500 == 0 or i == total_paras:
            pct = int(i / total_paras * 100)
            bar = "█" * (pct // 5) + "░" * (20 - pct // 5)
            print(f"\r    [{bar}] {pct}% ({i}/{total_paras} פסקאות)", end="", flush=True)
        if para.text.strip():
            text_parts.append(para.text)
    print()
    return "\n".join(text_parts)


def extract_text_from_pdf(filepath):
    reader = PdfReader(filepath)
    text_parts = []
    total_pages = len(reader.pages)
    for i, page in enumerate(reader.pages, 1):
        if i % 50 == 0 or i == total_pages:
            pct = int(i / total_pages * 100)
            bar = "█" * (pct // 5) + "░" * (20 - pct // 5)
            print(f"\r    [{bar}] {pct}% ({i}/{total_pages} עמודים)", end="", flush=True)
        page_text = page.extract_text()
        if page_text:
            text_parts.append(f"--- עמוד {i} ---\n{page_text}")
    print()
    return "\n\n".join(text_parts)


# ══════════════════════════════════════════════════════════════════════════════
#  Chunking (identical logic to RAGIndex in the original app.py)
# ══════════════════════════════════════════════════════════════════════════════

def split_into_chunks(text, min_chunk=200, max_chunk=2000, overlap_chars=150):
    raw_paragraphs = re.split(r"\n{2,}", text)
    chunks = []
    current = ""

    for para in raw_paragraphs:
        para = para.strip()
        if not para:
            continue
        is_section_start = bool(SECTION_PATTERNS.match(para))
        if current and (
            len(current) + len(para) + 1 > max_chunk
            or (is_section_start and len(current) >= min_chunk)
        ):
            chunks.append(current)
            current = para
        else:
            current = current + "\n" + para if current else para

    if current and len(current.strip()) > 30:
        chunks.append(current)

    # Merge tiny chunks
    merged = []
    for chunk in chunks:
        if merged and len(merged[-1]) < min_chunk:
            merged[-1] = merged[-1] + "\n" + chunk
        else:
            merged.append(chunk)

    # Overlap
    if overlap_chars > 0 and len(merged) > 1:
        overlapped = [merged[0]]
        for i in range(1, len(merged)):
            prev_tail = merged[i - 1][-overlap_chars:]
            overlapped.append(prev_tail + "\n" + merged[i])
        return overlapped

    return merged


# ══════════════════════════════════════════════════════════════════════════════
#  Daily-halacha snippet builder
# ══════════════════════════════════════════════════════════════════════════════

def build_daily_snippets(document_text):
    raw_paragraphs = re.split(r'\n{2,}', document_text)
    candidates = []
    last_topic = ""
    last_siman = ""

    for para in raw_paragraphs:
        para = para.strip()
        if not para:
            continue

        topic_match = _TOPIC_RE.match(para)
        if topic_match:
            last_topic = "הלכות " + topic_match.group(1)

        siman_match = _SIMAN_RE.match(para)
        if siman_match:
            last_siman = "סימן " + siman_match.group(1)

        if len(para) < 80:
            continue
        if not _CLEAN_START_RE.match(para):
            continue

        cleaned = re.sub(r'---\s*עמוד\s+\d+\s*---', '', para).strip()
        if len(cleaned) < 80:
            continue

        snippet_text = None
        if len(cleaned) <= 600:
            snippet_text = cleaned
        else:
            window = cleaned[:600]
            best_cut = -1
            for m in _SENTENCE_END_RE.finditer(window):
                best_cut = m.end()
            if best_cut >= 100:
                snippet_text = window[:best_cut].strip()
            else:
                for ch in ('.', ':', '׃'):
                    idx = window.rfind(ch)
                    if idx >= 100:
                        snippet_text = window[:idx + 1].strip()
                        break

        if snippet_text:
            candidates.append({
                "text": snippet_text,
                "topic": last_topic,
                "siman": last_siman,
            })

    return candidates


# ══════════════════════════════════════════════════════════════════════════════
#  Main build
# ══════════════════════════════════════════════════════════════════════════════

def main():
    os.makedirs(DATA_DIR, exist_ok=True)

    # ── 1. Load document ──────────────────────────────────────────────────
    docx_path = os.path.join(BASE_DIR, "ילקוט-יוסף-קיצור-שולחן-ערוך.docx")
    pdf_path = os.path.join(BASE_DIR, "ילקוט-יוסף-קיצור-שולחן-ערוך.pdf")

    if os.path.exists(docx_path):
        print(f"\n{'='*60}")
        print(f"[*] טוען מסמך DOCX: {os.path.basename(docx_path)}")
        print(f"{'='*60}")
        document_text = extract_text_from_docx(docx_path)
        document_filename = "ילקוט-יוסף-קיצור-שולחן-ערוך.docx"
    elif os.path.exists(pdf_path):
        print(f"\n{'='*60}")
        print(f"[*] טוען מסמך PDF: {os.path.basename(pdf_path)}")
        print(f"{'='*60}")
        document_text = extract_text_from_pdf(pdf_path)
        document_filename = "ילקוט-יוסף-קיצור-שולחן-ערוך.pdf"
    else:
        print("[!] שגיאה: לא נמצא קובץ ילקוט יוסף!")
        return

    char_count = len(document_text)
    page_count = document_text.count("--- עמוד") if "--- עמוד" in document_text else None

    print(f"[✓] מסמך נטען: {char_count:,} תווים")
    if page_count:
        print(f"    עמודים: {page_count}")

    # ── 2. Chunk ──────────────────────────────────────────────────────────
    print("\n[*] שלב 1/4: פיצול לקטעים...")
    chunks = split_into_chunks(document_text, min_chunk=200, max_chunk=2000, overlap_chars=150)
    print(f"    קטעים שנוצרו: {len(chunks):,}")

    # ── 3. Build word-level TF-IDF ────────────────────────────────────────
    print("[*] שלב 2/4: בניית אינדקס מילים (TF-IDF)...")
    t0 = time.time()
    word_vectorizer = TfidfVectorizer(
        analyzer="word",
        token_pattern=r"[\u0590-\u05FF\w]{2,}",
        sublinear_tf=True,
        min_df=2,
        max_df=0.85,
        ngram_range=(1, 2),
        max_features=5000,
        dtype=np.float32,
    )
    word_matrix = word_vectorizer.fit_transform(chunks)
    print(f"    מילון מילים: {len(word_vectorizer.vocabulary_):,} מונחים")

    # ── 4. Build char n-gram TF-IDF ───────────────────────────────────────
    print("[*] שלב 3/4: בניית אינדקס תווים (fuzzy matching)...")
    char_vectorizer = TfidfVectorizer(
        analyzer="char_wb",
        ngram_range=(3, 5),
        min_df=3,
        max_df=0.9,
        sublinear_tf=True,
        max_features=8000,
        dtype=np.float32,
    )
    char_matrix = char_vectorizer.fit_transform(chunks)
    print(f"    מילון תווים: {len(char_vectorizer.vocabulary_):,} n-grams")
    elapsed = time.time() - t0
    print(f"    זמן בניית אינדקס: {elapsed:.1f} שניות")

    # ── 5. Build daily snippets ───────────────────────────────────────────
    print("[*] שלב 4/4: בניית קטעי הלכה יומית...")
    snippets = build_daily_snippets(document_text)
    print(f"    קטעים ראויים: {len(snippets):,}")

    # ── 6. Save everything to data/ ───────────────────────────────────────
    print(f"\n[*] שומר לתיקיית {DATA_DIR}...")

    joblib.dump(chunks, os.path.join(DATA_DIR, "chunks.joblib"))
    joblib.dump(word_vectorizer, os.path.join(DATA_DIR, "word_vectorizer.joblib"))
    joblib.dump(word_matrix, os.path.join(DATA_DIR, "word_matrix.joblib"))
    joblib.dump(char_vectorizer, os.path.join(DATA_DIR, "char_vectorizer.joblib"))
    joblib.dump(char_matrix, os.path.join(DATA_DIR, "char_matrix.joblib"))
    joblib.dump(snippets, os.path.join(DATA_DIR, "daily_snippets.joblib"))

    # Save metadata as JSON for easy inspection
    metadata = {
        "document_filename": document_filename,
        "char_count": char_count,
        "page_count": page_count,
        "chunk_count": len(chunks),
        "snippet_count": len(snippets),
        "word_vocab_size": len(word_vectorizer.vocabulary_),
        "char_vocab_size": len(char_vectorizer.vocabulary_),
    }
    with open(os.path.join(DATA_DIR, "metadata.json"), 'w', encoding='utf-8') as f:
        json.dump(metadata, f, ensure_ascii=False, indent=2)

    # Print file sizes
    print("\n[✓] קבצים שנשמרו:")
    total_size = 0
    for fname in os.listdir(DATA_DIR):
        fpath = os.path.join(DATA_DIR, fname)
        size = os.path.getsize(fpath)
        total_size += size
        print(f"    {fname}: {size / 1024:.1f} KB")
    print(f"    ────────────────────────")
    print(f"    סה\"כ: {total_size / 1024:.1f} KB ({total_size / (1024*1024):.1f} MB)")

    print(f"\n{'='*60}")
    print("[✓] הבנייה הושלמה! כעת ניתן לפרוס ל-Render.")
    print(f"{'='*60}\n")


if __name__ == "__main__":
    main()
